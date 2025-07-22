import requests
import json
import docx
import pypdf
import os
import pandas as pd
from tabulate import tabulate
import signal
from contextlib import contextmanager
import re

# LlamaIndex core imports for persistence
from llama_index.core import StorageContext, load_index_from_storage

# LlamaIndex and Unstructured imports
from llama_index.llms.ollama import Ollama
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core import Settings
from llama_index.core import VectorStoreIndex, Document
from llama_index.core.response_synthesizers import CompactAndRefine
from llama_index.core.prompts import PromptTemplate

# NEW IMPORTS FOR UNSTRUCTURED NODE PARSING
from llama_index.core.node_parser import UnstructuredElementNodeParser
from unstructured.partition.pdf import partition_pdf
from unstructured.partition.docx import partition_docx
from unstructured.documents.elements import Table, Text # Import specific element types
# Import the ElementMetadata class to check its type if needed
from unstructured.documents.elements import ElementMetadata


def extract_text_from_pdf_simple(file_path):
    """
    Extracts text from a PDF file page by page.
    Note: This is a simple text extraction and might not preserve complex layouts
    or table structures as effectively as 'unstructured'.
    """
    text = ""
    try:
        reader = pypdf.PdfReader(file_path)
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text += page.extract_text() + "\n" # Add newline between pages
        return text
    except FileNotFoundError:
        print(f"Error: PDF file not found at {file_path}")
        return None
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None


def extract_text_from_docx(file_path):
    """Extracts text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except FileNotFoundError:
        print(f"Error: DOCX file not found at {file_path}")
        return None
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None


# --- MODIFIED: read_document_content to return LlamaIndex Documents for nodes ---
def read_document_and_create_nodes(file_path):
    """
    Reads text and tables from a PDF or DOCX file using unstructured,
    and returns a list of LlamaIndex Document objects.
    """
    elements = []
    try:
        if file_path.lower().endswith('.docx'):
            print("Using unstructured.partition_docx for DOCX file...")
            elements = partition_docx(
                filename=file_path,
                infer_table_structure=True,
            )
        elif file_path.lower().endswith('.pdf'):
            print("Using unstructured.partition_pdf for PDF file...")
            elements = partition_pdf(
                filename=file_path,
                infer_table_structure=True,
                strategy="hi_res",
                # Add 'size' parameter as suggested by unstructured's deprecation warning
                # 'longest_edge' is a common and usually good default
                # For example, if you want to resize images to a max of 2000 pixels on their longest edge
                # You can adjust this value as needed.
                # If you don't extract image blocks, you can remove these lines related to size.
                extract_image_block_types=["Image"],
                extract_image_block_to_payload=True,
                # size={"longest_edge": 2000}, # Example to address the deprecation warning
                chunking_strategy="by_title",
                max_characters=10000,
                combine_text_under_n_chars=2000,
                new_after_n_chars=6000,
            )
        else:
            print("Unsupported file format. Please provide a .pdf or .docx file.")
            return []

        # Create LlamaIndex Documents from Unstructured elements
        llama_documents = []
        for el in elements:
            # Safely get page_number from metadata
            page_number = None
            if hasattr(el, 'metadata') and el.metadata is not None:
                # Access attributes directly from ElementMetadata object
                page_number = getattr(el.metadata, 'page_number', 'N/A')

            # Unstructured Table elements have a .text_as_html or .text_as_markdown
            if isinstance(el, Table):
                # Prefer markdown for LLMs as it's often cleaner for table structure
                # Check if text_as_html or text_as_markdown exists, otherwise use plain text
                table_content = el.text
                if hasattr(el, 'text_as_html') and el.text_as_html:
                    table_content = el.text_as_html
                elif hasattr(el, 'text_as_markdown') and el.text_as_markdown:
                    table_content = el.text_as_markdown

                # Add metadata to distinguish table nodes
                llama_documents.append(Document(
                    text=table_content,
                    metadata={
                        "element_type": "table",
                        "original_unstructured_text": el.text, # Keep original text as well
                        "page_number": page_number
                    }
                ))
            elif isinstance(el, Text): # Standard text elements
                llama_documents.append(Document(
                    text=el.text,
                    metadata={
                        "element_type": "text",
                        "page_number": page_number
                    }
                ))
            else: # Catch-all for other element types (e.g., Title, NarrativeText, etc.)
                if el.text and el.text.strip(): # Only add if there's actual text
                    llama_documents.append(Document(
                        text=el.text,
                        metadata={
                            "element_type": getattr(el, 'category', 'other_element').lower(),
                            "page_number": page_number
                        }
                    ))
        
        return llama_documents

    except FileNotFoundError:
        print(f"Error: File not found at {file_path}")
        return []
    except Exception as e:
        print(f"Error extracting content from {file_path}: {e}")
        return []


import pandas as pd
from tabulate import tabulate # Assuming tabulate is needed for display_results


def json_to_dataframe(json_data, characteristic_name):
    """
    Convert JSON data to pandas DataFrame with proper formatting for table display.
    Ensures treatments are always columns and metrics are always rows.
    """
    if not isinstance(json_data, dict):
        return None

    # Check the structure to determine the correct orientation
    # Handle empty dictionary for characteristic_data
    if not json_data:
        return pd.DataFrame() # Return an empty DataFrame

    sample_key = next(iter(json_data.keys()))
    sample_value = json_data[sample_key]


    if isinstance(sample_value, dict):
        # Check if the nested values contain treatment names
        nested_keys = list(sample_value.keys())
        treatment_names_and_total = ['Treatment A', 'Treatment B', 'Total']


        # If nested keys are treatment names, this is Format 2 (correct structure for stratification/categorical)
        if any(treatment in nested_keys for treatment in treatment_names_and_total):
            # --- MODIFICATION START ---
            # Instead of just transposing, we'll explicitly create the DataFrame
            # to ensure column order and presence of 'Total'.


            # Determine all possible columns (treatments and Total, in desired order)
            all_found_columns = set()
            for level_data in json_data.values():
                if isinstance(level_data, dict):
                    all_found_columns.update(level_data.keys())


            ordered_columns = []
            for col in treatment_names_and_total: # Prioritize 'Treatment A', 'Treatment B', 'Total'
                if col in all_found_columns:
                    ordered_columns.append(col)
                    all_found_columns.remove(col) # Remove to avoid duplicates


            # Add any other columns found that were not in our prioritized list
            ordered_columns.extend(sorted(list(all_found_columns)))


            df_data = {}
            for metric, values in json_data.items():
                df_data[metric] = {col: values.get(col, '') for col in ordered_columns}


            df = pd.DataFrame.from_dict(df_data, orient='index')
            df.index.name = "Metric" # Or "Category" depending on context, "Metric" for general
            # --- MODIFICATION END ---
            return df


        # If nested keys are metrics, this is Format 1 (needs restructuring)
        else:
            # Format 1: {"Treatment A": {"n": "xx", "Mean": "xxx.xx", ...}}
            # We need to restructure this to have metrics as rows and treatments as columns


            # First, check if the outer keys are treatment names
            outer_keys = list(json_data.keys())
            if any(treatment in outer_keys for treatment in treatment_names_and_total):
                # Restructure: transpose the nested structure
                restructured_data = {}


                # Get all possible metrics from all treatments
                all_metrics = set()
                for treatment_data in json_data.values():
                    if isinstance(treatment_data, dict):
                        all_metrics.update(treatment_data.keys())


                # Build the restructured data with metrics as outer keys
                for metric in all_metrics:
                    restructured_data[metric] = {}
                    for treatment, treatment_data in json_data.items():
                        if isinstance(treatment_data, dict) and metric in treatment_data:
                            restructured_data[metric][treatment] = treatment_data[metric]
                        else:
                            restructured_data[metric][treatment] = ""  # or None


                df = pd.DataFrame(restructured_data).T
                df.index.name = "Metric"
                return df


            else:
                # This might be categorical data that is already correctly structured for transpose
                # e.g., {"Category A": {"Subcategory 1": "val"}}
                df = pd.DataFrame(json_data).T
                df.index.name = "Category" # Fallback for general categorical
                return df


    # This handles cases where the immediate values are not dictionaries,
    # e.g., {"Key": "Value"}
    return None # Or handle more specifically if there's a different expected output for this case


def clean_llm_response(response_text):
    """
    Clean the LLM response to fix common JSON formatting issues.
    """
    # Remove markdown code blocks if present
    response_text = response_text.strip()
    if response_text.startswith("```json") and response_text.endswith("```"):
        response_text = response_text[len("```json"):-len("```")].strip()
    elif response_text.startswith("```") and response_text.endswith("```"):
        response_text = response_text[3:-3].strip()
    
    # Fix double curly braces to single curly braces
    # This handles the template escaping issue
    response_text = response_text.replace("{{", "{").replace("}}", "}")

    # Remove single-line comments (//) and any lines that look like comments or instructions
    response_text = re.sub(r'^\s*\/\/.*$', '', response_text, flags=re.MULTILINE)
    
    # Aggressively remove common LLM truncation patterns *within* or at the end of JSON
    response_text = re.sub(r',\s*\.\.\.', '', response_text) # Remove ", ..."
    response_text = re.sub(r'\s*\.\.\.', '', response_text) # Remove "..." if not preceded by comma
    
    # Attempt to insert missing commas where a closing brace is immediately followed by a new key
    # This targets patterns like '} "NewKey"' and changes them to '}, "NewKey"'
    response_text = re.sub(r'}\s*"([^"\s]+)"\s*:', r'}, "\1":', response_text)
    
    # IMPROVED: Extract only the JSON part - remove any trailing text after the JSON
    # Look for the last closing brace that matches the first opening brace
    try:
        # Find the first opening brace
        start_idx = response_text.find('{')
        if start_idx == -1:
            return response_text
            
        # Count braces to find the matching closing brace
        brace_count = 0
        end_idx = start_idx
        
        for i, char in enumerate(response_text[start_idx:], start_idx):
            if char == '{':
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0:
                    end_idx = i
                    break
        
        # Extract only the JSON part
        if end_idx > start_idx:
            response_text = response_text[start_idx:end_idx + 1]
    except Exception as e:
        print(f"Warning: Could not extract JSON cleanly: {e}")
    
    return response_text


def display_results(parsed_json, query_text):
    """
    Display the parsed JSON results as formatted tables.
    """
    if not isinstance(parsed_json, dict) or len(parsed_json) == 0:
        print("No data found or empty response.")
        return False
    
    processed_any = False
    
    for characteristic_name, characteristic_data in parsed_json.items():
        if isinstance(characteristic_data, dict):
            # Try to convert to DataFrame
            df = json_to_dataframe(characteristic_data, characteristic_name)
            
            if df is not None and not df.empty:
                print(f"\n--- {characteristic_name} ---")
                print(tabulate(df, headers='keys', tablefmt='grid', showindex=True))
                print()
                processed_any = True
            else:
                # Fallback: display as formatted JSON
                print(f"\n--- {characteristic_name} (JSON Format) ---")
                print(json.dumps({characteristic_name: characteristic_data}, indent=2))
                print()
                processed_any = True
        
        elif isinstance(characteristic_data, (str, int, float, bool, type(None))):
            # Handle simple key-value pairs
            print(f"\n--- Specific Value ---")
            print(f"{characteristic_name}: {characteristic_data}")
            print()
            processed_any = True
    
    if not processed_any:
        print("Could not process the JSON into table format. Raw JSON:")
        print(json.dumps(parsed_json, indent=2))
    
    return processed_any


Settings.embed_model = HuggingFaceEmbedding(model_name="BAAI/bge-base-en-v1.5")
Settings.llm = Ollama(model="mistral", request_timeout=1000.0, format="json")


@contextmanager
def timeout_context(seconds):
    """Context manager for timeout handling"""
    def timeout_handler(signum, frame):
        raise TimeoutError(f"Operation timed out after {seconds} seconds")

    signal.signal(signal.SIGALRM, timeout_handler)
    signal.alarm(seconds)
    try:
        yield
    finally:
        signal.alarm(0)


if __name__ == '__main__':
    base_directory_path = "/Users/Sravya/Desktop/AI_model_table_shells"
    file_name = "Table shell standard.docx" # Make sure this matches your file
    file_path = os.path.join(base_directory_path, file_name)


    # --- Index Storage Setup ---
    PERSIST_DIR = "./storage" # Define the directory to store the index


    index = None # Initialize index variable


    if not os.path.exists(file_path):
        print(f"Error: Document file not found at {file_path}. Please ensure the file exists and the 'file_name' variable is correct.")
        exit() # Exit if source document not found


    # Check if the index already exists in storage
    if not os.path.exists(PERSIST_DIR) or not os.listdir(PERSIST_DIR):
        print("Index not found in storage. Creating a new index...")
        
        # --- MODIFIED: Call the new function to get LlamaIndex Documents from unstructured elements ---
        llama_documents = read_document_and_create_nodes(file_path)

        if not llama_documents:
            print("No content extracted from the document or no valid nodes created. Exiting RAG pipeline setup.")
            exit() # Exit if no documents were loaded for indexing

        print(f"\n--- First 5 LlamaIndex Documents Created (partial view) ---")
        for i, doc in enumerate(llama_documents[:5]):
            print(f"Doc {i+1} (type: {doc.metadata.get('element_type', 'unknown')}, page: {doc.metadata.get('page_number', 'N/A')}): {doc.text[:200]}..." if doc.text else "Empty Document")
        print("--------------------------------------------------\n")

        print("Creating VectorStoreIndex (this involves embedding the document nodes)...")
        # --- MODIFIED: Create NodeParser and then index the parsed nodes ---
        # The UnstructuredElementNodeParser will split documents into smaller chunks (nodes)
        # and handle tables specifically.
        node_parser = UnstructuredElementNodeParser(
            # You can customize chunking parameters here if needed,
            # but Unstructured's own chunking is often a good starting point.
            # max_characters=1000,
            # combine_text_under_n_chars=500,
            # new_after_n_chars=1500,
        )
        nodes = node_parser.get_nodes_from_documents(llama_documents)

        print(f"Generated {len(nodes)} LlamaIndex Nodes for embedding.")
        for i, node in enumerate(nodes[:5]):
            print(f"Node {i+1} (type: {node.metadata.get('element_type', 'unknown')}, page: {node.metadata.get('page_number', 'N/A')}): {node.text[:200]}..." if node.text else "Empty Node")


        index = VectorStoreIndex(nodes, embed_model=Settings.embed_model) # Index the generated nodes
        print("VectorStoreIndex created. Persisting index to disk...")
        index.storage_context.persist(persist_dir=PERSIST_DIR) # Persist the index
        print(f"Index persisted to {PERSIST_DIR}")
    else:
        print(f"Loading index from {PERSIST_DIR}...")
        storage_context = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
        index = load_index_from_storage(storage_context) # Load the index from storage
        print("Index loaded.")


    print("Creating QueryEngine...")


    # FIXED PROMPT TEMPLATE - This is the key change
    custom_qa_template_str = """Context information is below.
---------------------
{context_str}
---------------------
Given the context information above, extract the table shell structure for the requested characteristic.

IMPORTANT INSTRUCTIONS:
1. This is a TABLE SHELL document with PLACEHOLDER values that should be preserved exactly as they appear
2. Extract the EXACT structure and placeholder values from the document (like "xx", "xxx.xx", "xx (xx.x)", etc.)
3. DO NOT replace placeholders with empty strings - keep them as they appear in the document
4. Your response MUST be valid JSON ONLY
5. Do NOT include any conversational text, explanations, or markdown

The document contains table shells with these placeholder patterns:
- "xx" for count placeholders
- "xxx.xx" for decimal placeholders  
- "xxx.xxx" for 3-decimal placeholders
- "xx (xx.x)" for count with percentage placeholders
- "xxx, xxx" for min/max range placeholders

For the query "{query_str}", find the corresponding table structure and return it with the EXACT placeholders as shown in the document.

EXAMPLE OUTPUT STRUCTURE:

For continuous variables:
{{
  "Height (cm)": {{
    "n": {{
      "Treatment A": "xx",
      "Treatment B": "xx", 
      "Total": "xx"
    }},
    "Mean": {{
      "Treatment A": "xxx.xx",
      "Treatment B": "xxx.xx",
      "Total": "xxx.xx"
    }},
    "Standard Deviation": {{
      "Treatment A": "xxx.xxx",
      "Treatment B": "xxx.xxx", 
      "Total": "xxx.xxx"
    }},
    "Median": {{
      "Treatment A": "xxx.xx",
      "Treatment B": "xxx.xx",
      "Total": "xxx.xx"
    }},
    "Min, Max": {{
      "Treatment A": "xxx, xxx",
      "Treatment B": "xxx, xxx",
      "Total": "xxx, xxx"
    }}
  }}
}}

For categorical variables:
{{
  "Ethnicity [for US studies only]": {{
    "Hispanic/Latino": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "Non-Hispanic/Non-Latino": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "Unknown": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)", 
      "Total": "xx (xx.x)"
    }}
  }}
}}

For Stratification Factors:
{{
  "Stratification Factor Name": {{
    "Level 1": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "Level 2": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }}
  }}
}}

For 12-Lead ECG:
{{
  "12-Lead ECG": {{
    "Normal": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "Abnormal, Not Clinically Significant": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "Clinically significant findings": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "Not Performed": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }}
  }}
}}

For Baseline/Biomarker Subgroups:
{{
  "Baseline/Biomarker Subgroups": {{
    "Category 1": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "Category 2": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }}
  }}
}}

For Bone marrow/aspirate blast count at baseline:
{{
  "Height (cm)": {{
    "n": {{
      "Treatment A": "xx",
      "Treatment B": "xx", 
      "Total": "xx"
    }},
    "Mean": {{
      "Treatment A": "xxx.xx",
      "Treatment B": "xxx.xx",
      "Total": "xxx.xx"
    }},
    "Standard Deviation": {{
      "Treatment A": "xxx.xxx",
      "Treatment B": "xxx.xxx", 
      "Total": "xxx.xxx"
    }},
    "Median": {{
      "Treatment A": "xxx.xx",
      "Treatment B": "xxx.xx",
      "Total": "xxx.xx"
    }},
    "Min, Max": {{
      "Treatment A": "xxx, xxx",
      "Treatment B": "xxx, xxx",
      "Total": "xxx, xxx"
    }}
    "<Median": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }}
    ">=Median": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }}
  }}
}}

For Region/Country of Enrollment:
{{
  "Region/Country of Enrollment": {{
    "North America": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "   Country 1": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "Asia": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "   Country 1": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "Europe (required to list all countries)": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "   Country 1": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }}
    "   Country 2": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }}
    "   Country n": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }}
    "Rest of World": {{
      "Treatment A": "xx (xx.x)",
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }},
    "   Country 1": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }}
    "   Country 2": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }}
    "   Country n": {{
      "Treatment A": "xx (xx.x)", 
      "Treatment B": "xx (xx.x)",
      "Total": "xx (xx.x)"
    }}
  }}
}}

For Time from Initial Histologic Diagnosis to Randomization (days):
{{
  "Time from Initial Histologic Diagnosis to Randomization (days)": {{
    "n": {{
      "Treatment A": "xx",
      "Treatment B": "xx", 
      "Total": "xx"
    }},
    "Mean": {{
      "Treatment A": "xxx.xx",
      "Treatment B": "xxx.xx",
      "Total": "xxx.xx"
    }},
    "Standard Deviation": {{
      "Treatment A": "xxx.xxx",
      "Treatment B": "xxx.xxx", 
      "Total": "xxx.xxx"
    }},
    "Median": {{
      "Treatment A": "xxx.xx",
      "Treatment B": "xxx.xx",
      "Total": "xxx.xx"
    }},
    "Min, Max": {{
      "Treatment A": "xxx, xxx",
      "Treatment B": "xxx, xxx",
      "Total": "xxx, xxx"
    }}
  }}
}}

Return ONLY the JSON for the requested table structure.

Query: {query_str}
"""
    custom_qa_template = PromptTemplate(custom_qa_template_str)


    query_engine = index.as_query_engine(
        llm=Settings.llm,
        response_synthesizer=CompactAndRefine(text_qa_template=custom_qa_template),
        similarity_top_k=3  # INCREASED to get better context coverage
    )
    print("QueryEngine created.")


    # Updated queries for better matching
    queries = [
        "Height table structure",
        "Weight table structure",
        "Ethnicity [for US studies only] table structure",
        "Body Mass Index (kg/m2) table structure",
        "Body Surface Area (m2) table structure",
        "Stratification Factor 1 (EDC) table structure",
        "Stratification Factor 1 (IXRS) table structure",
        "ECOG Performance Status table structure",
        "12-Lead ECG table structure",
        "Baseline/Biomarker Subgroups table structure",
        "Smoking Status table structure",
        "Bone marrow/aspirate blast count at baseline table structure",

    ]


    print("\n--- Running Queries with RAG ---")
    for i, query_text in enumerate(queries):
        print(f"\n{'='*60}")
        print(f"Query {i+1}: {query_text}")
        print('='*60)


        try:
            with timeout_context(600):
                response = query_engine.query(query_text)
                llm_response_text = response.response


            # Clean up the response
            llm_response_text = clean_llm_response(llm_response_text)
            
            print(f"DEBUG - Cleaned response: {llm_response_text[:500]}...")  # Debug output


            try:
                # Parse the JSON response
                parsed_json = json.loads(llm_response_text)
                
                # Display the results as formatted tables
                success = display_results(parsed_json, query_text)
                
                if not success:
                    print("Failed to process response into table format.")


            except json.JSONDecodeError as e:
                print(f"Could not parse as JSON. Raw response:")
                print(f"Error: {e}")
                print(f"Response: {llm_response_text}")


        except TimeoutError as te:
            print(f"Query timed out: {te}")
        except Exception as e:
            print(f"Error during query {i+1}: {e}")
            
    print(f"\n{'='*60}")
    print("All queries completed.")
    print('='*60)